"""
CV / Cover Letter Processor.

How it works:
1. User provides a DOCX template where the sections they want AI to fill
   are highlighted in a specific colour (default: YELLOW).
2. This module parses the DOCX, extracts those coloured sections,
   sends them to the AI generator, then replaces them in-place.
3. Only the coloured runs are touched — ALL other formatting is preserved.

Colour detection:
  - Highlight (background): checked via run.font.highlight_color
  - Font colour (foreground): checked via run.font.color.rgb
  Configurable via CV_HIGHLIGHT_COLOR in .env
"""
from __future__ import annotations

import copy
import re
import uuid
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.shared import RGBColor

import config
from utils.logger import logger

# Mapping of colour names to WD_COLOR_INDEX values
HIGHLIGHT_MAP: Dict[str, WD_COLOR_INDEX] = {
    "YELLOW": WD_COLOR_INDEX.YELLOW,
    "CYAN": WD_COLOR_INDEX.CYAN,
    "GREEN": WD_COLOR_INDEX.GREEN,
    "PINK": WD_COLOR_INDEX.PINK,
    "RED": WD_COLOR_INDEX.RED,
    "BLUE": WD_COLOR_INDEX.BLUE,
    "TURQUOISE": WD_COLOR_INDEX.TURQUOISE,
    "DARK_YELLOW": WD_COLOR_INDEX.DARK_YELLOW,
    "VIOLET": WD_COLOR_INDEX.VIOLET,
}

# Mapping for font colour detection (foreground)
FONT_COLOR_MAP: Dict[str, RGBColor] = {
    "RED": RGBColor(0xFF, 0x00, 0x00),
    "BLUE": RGBColor(0x00, 0x00, 0xFF),
    "GREEN": RGBColor(0x00, 0xFF, 0x00),
    "YELLOW": RGBColor(0xFF, 0xFF, 0x00),
    "CYAN": RGBColor(0x00, 0xFF, 0xFF),
    "PINK": RGBColor(0xFF, 0x00, 0xFF),
}


class CVProcessor:
    """Parse, extract and replace coloured sections in a DOCX template."""

    def __init__(self, colour: Optional[str] = None):
        self.colour = (colour or config.CV_HIGHLIGHT_COLOR).upper()
        self.highlight_value: Optional[WD_COLOR_INDEX] = HIGHLIGHT_MAP.get(self.colour)
        self.font_color_value: Optional[RGBColor] = FONT_COLOR_MAP.get(self.colour)

    # ── Public API ─────────────────────────────────────────────

    def extract_sections(self, docx_path: Path) -> List[Dict]:
        """
        Scan the DOCX and return a list of section descriptors:
        [
          {"id": "sec_abc123", "placeholder": "Current text in this coloured run group"},
          ...
        ]
        """
        doc = Document(str(docx_path))
        sections = []

        for para in doc.paragraphs:
            sections.extend(self._extract_from_paragraph(para))

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        sections.extend(self._extract_from_paragraph(para))

        logger.info(f"Extracted {len(sections)} modifiable sections from {docx_path.name}")
        return sections

    def apply_replacements(
        self,
        docx_path: Path,
        replacements: Dict[str, str],
        output_path: Path,
    ) -> Path:
        """
        Write a new DOCX at output_path with all coloured runs replaced
        by the generated content in `replacements`.
        """
        doc = Document(str(docx_path))

        for para in doc.paragraphs:
            self._replace_in_paragraph(para, replacements)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_in_paragraph(para, replacements)

        doc.save(str(output_path))
        logger.info(f"Saved modified document to {output_path}")
        return output_path

    # ── Internal Helpers ───────────────────────────────────────

    def _is_colored_run(self, run) -> bool:
        """Return True if this run is marked with the target colour."""
        # Check highlight (background colour)
        if self.highlight_value is not None:
            try:
                if run.font.highlight_color == self.highlight_value:
                    return True
            except Exception:
                pass

        # Check font (foreground) colour
        if self.font_color_value is not None:
            try:
                if run.font.color.type and run.font.color.rgb == self.font_color_value:
                    return True
            except Exception:
                pass

        return False

    def _extract_from_paragraph(self, para) -> List[Dict]:
        """
        Group consecutive coloured runs within a paragraph into sections.
        Each contiguous group of coloured runs becomes one section.
        """
        sections = []
        buffer_text = ""
        buffer_start = None

        for run in para.runs:
            if self._is_colored_run(run):
                if buffer_start is None:
                    buffer_start = run
                buffer_text += run.text
            else:
                if buffer_text.strip():
                    sec_id = f"sec_{uuid.uuid4().hex[:8]}"
                    sections.append({
                        "id": sec_id,
                        "placeholder": buffer_text,
                        "_para_id": id(para),
                    })
                buffer_text = ""
                buffer_start = None

        # Flush final buffer
        if buffer_text.strip():
            sec_id = f"sec_{uuid.uuid4().hex[:8]}"
            sections.append({
                "id": sec_id,
                "placeholder": buffer_text,
                "_para_id": id(para),
            })

        return sections

    def _replace_in_paragraph(self, para, replacements: Dict[str, str]) -> None:
        """
        Replace the text of coloured runs with the generated content.
        Strategy: collect consecutive coloured runs, merge them, write
        the new text into the first run, clear the rest.
        """
        colored_groups: List[List[int]] = []
        current_group: List[int] = []

        for i, run in enumerate(para.runs):
            if self._is_colored_run(run):
                current_group.append(i)
            else:
                if current_group:
                    colored_groups.append(current_group)
                    current_group = []
        if current_group:
            colored_groups.append(current_group)

        if not colored_groups:
            return

        # We match groups to replacements by order (same order as extraction)
        # Build a reverse lookup: original text → replacement
        all_colored_text = "".join(
            para.runs[i].text for group in colored_groups for i in group
        )

        # Find a matching replacement by placeholder similarity
        best_replacement: Optional[str] = None
        best_score = 0
        for rep_placeholder, rep_text in replacements.items():
            # Simple: use first replacement that isn't consumed yet
            # (More sophisticated: fuzzy match placeholder text)
            best_replacement = rep_text
            break

        if best_replacement is None:
            return

        # Write replacement into first run of first group, clear others
        first_group = colored_groups[0]
        if first_group:
            first_run = para.runs[first_group[0]]
            first_run.text = best_replacement
            # Remove highlight from replaced run so it looks normal
            try:
                first_run.font.highlight_color = None
            except Exception:
                pass
            # Clear remaining runs in this group
            for idx in first_group[1:]:
                para.runs[idx].text = ""

    def extract_sections_ordered(self, docx_path: Path) -> Tuple[List[Dict], List[Tuple]]:
        """
        Returns sections in document order together with their paragraph/run context
        so we can do precise ordered replacement.
        """
        doc = Document(str(docx_path))
        sections = []
        contexts = []  # (para_ref, [run_indices])

        def process_para(para):
            current_group_runs = []
            buffer_text = ""

            for i, run in enumerate(para.runs):
                if self._is_colored_run(run):
                    current_group_runs.append(i)
                    buffer_text += run.text
                else:
                    if buffer_text.strip():
                        sec_id = f"sec_{uuid.uuid4().hex[:8]}"
                        sections.append({"id": sec_id, "placeholder": buffer_text})
                        contexts.append((para, list(current_group_runs)))
                    buffer_text = ""
                    current_group_runs = []

            if buffer_text.strip():
                sec_id = f"sec_{uuid.uuid4().hex[:8]}"
                sections.append({"id": sec_id, "placeholder": buffer_text})
                contexts.append((para, list(current_group_runs)))

        for para in doc.paragraphs:
            process_para(para)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        process_para(para)

        return sections, contexts

    def apply_ordered_replacements(
        self,
        docx_path: Path,
        replacements: Dict[str, str],  # section_id → new_text
        output_path: Path,
    ) -> Path:
        """
        Precise ordered replacement. Matches sections by their ID
        (from extract_sections_ordered) and writes the replacement text
        into exactly the correct runs.
        """
        doc = Document(str(docx_path))
        sections, contexts = self.extract_sections_ordered(docx_path)

        # Build id→replacement map
        id_to_text = {sec["id"]: replacements.get(sec["id"], sec["placeholder"]) for sec in sections}

        # Re-walk the document with a document that was freshly loaded for writing
        doc2 = Document(str(docx_path))

        def process_para_write(para, write_doc_para):
            section_idx = 0
            current_group_runs_idx = []
            buffer_text = ""
            pending_writes: List[Tuple[List[int], str]] = []

            for i, run in enumerate(para.runs):
                if self._is_colored_run(run):
                    current_group_runs_idx.append(i)
                    buffer_text += run.text
                else:
                    if buffer_text.strip() and section_idx < len(sections):
                        sec = sections[section_idx]
                        new_text = id_to_text.get(sec["id"], buffer_text)
                        pending_writes.append((list(current_group_runs_idx), new_text))
                        section_idx += 1
                    buffer_text = ""
                    current_group_runs_idx = []

            if buffer_text.strip() and section_idx < len(sections):
                sec = sections[section_idx]
                new_text = id_to_text.get(sec["id"], buffer_text)
                pending_writes.append((list(current_group_runs_idx), new_text))

            # Apply writes
            for run_indices, new_text in pending_writes:
                if run_indices:
                    write_doc_para.runs[run_indices[0]].text = new_text
                    try:
                        write_doc_para.runs[run_indices[0]].font.highlight_color = None
                    except Exception:
                        pass
                    for idx in run_indices[1:]:
                        write_doc_para.runs[idx].text = ""

        write_paras = list(doc2.paragraphs)
        source_paras = list(doc.paragraphs)

        for sp, wp in zip(source_paras, write_paras):
            process_para_write(sp, wp)

        doc2.save(str(output_path))
        logger.info(f"Saved ordered replacement document → {output_path}")
        return output_path
