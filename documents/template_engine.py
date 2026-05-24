"""
Template Engine — purpose-built for the CV/CL DOCX templates.

Template structure (detected from inspection):
  - Fixed paragraphs have font color #80340D (headings) or #153D63 (company/role/dates)
  - Variable paragraphs (to be replaced per job) have no special color

CV Variable Paragraph Indices (Diksha Desai template):
  summary        : 4
  competencies   : 6
  chintamani_1..4: 17,18,19,20  (Chintamani Thermal Technologies bullets)
  accenture_1..4 : 23,24,25,26  (Accenture Solutions bullets)
  project1_desc  : 29           (Supplier Spend Analytics — Objective line)
  project2_desc  : 34           (Insurance Operations — Objective line)

CL Variable Paragraph Indices:
  date_line      : 2  (tab + date portion only)
  company_name   : 3
  company_addr   : 4
  subject_line   : 5
  para1          : 6
  para2          : 7
  para3          : 8
  para4          : 9
  para5          : 10
"""
from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Tuple

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from utils.logger import logger

# ── Paragraph index maps ───────────────────────────────────────

CV_SECTIONS = {
    "summary":        [4],
    "competencies":   [6],
    "chintamani":     [17, 18, 19, 20],
    "accenture":      [23, 24, 25, 26],
    "project1_desc":  [29],
    "project2_desc":  [34],
}

CL_SECTIONS = {
    "company_name":    [3],
    "company_addr":    [4],
    "subject_line":    [5],
    "para1":           [6],
    "para2":           [7],
    "para3":           [8],
    "para4":           [9],
    "para5":           [10],
}


class TemplateEngine:
    """
    Replace variable paragraphs in the CV/CL DOCX templates with
    AI-generated content while preserving all fixed formatting.
    """

    # ── CV ──────────────────────────────────────────────────────

    def apply_cv_content(
        self,
        template_path: Path,
        content: Dict[str, str | List[str]],
        output_path: Path,
    ) -> Path:
        """
        Write tailored content into the CV template.

        Expected content keys:
          summary        : str   (plain text, ~60 words)
          competencies   : str   (comma-separated, ~60 words)
          chintamani     : list of 4 strings, each "Bold Label: description."
          accenture      : list of 4 strings, each "Bold Label: description."
          project1_desc  : str   (~50 words, Supplier Spend Analytics objective)
          project2_desc  : str   (~50 words, Insurance Operations objective)
        """
        doc = Document(str(template_path))
        paras = doc.paragraphs

        # Fix EMF icons before any content changes so they render in LibreOffice PDF
        self._fix_emf_icons(doc)

        for section, indices in CV_SECTIONS.items():
            value = content.get(section)
            if not value:
                logger.debug(f"CV: no content for section '{section}', skipping")
                continue

            if isinstance(value, list):
                # Bullet lists — one string per bullet
                for idx, text in zip(indices, value):
                    if idx < len(paras):
                        self._replace_bullet(paras[idx], text.strip())
            else:
                if indices[0] < len(paras):
                    self._replace_plain(paras[indices[0]], value.strip())

        # Update date on last paragraph (Name\tDate: ...)
        self._update_date(paras)

        doc.save(str(output_path))
        logger.info(f"CV saved → {output_path.name}")
        return output_path

    # ── CL ──────────────────────────────────────────────────────

    def apply_cl_content(
        self,
        template_path: Path,
        content: Dict[str, str],
        output_path: Path,
    ) -> Path:
        """
        Write tailored content into the CL template.

        Expected content keys:
          company_name  : str  (e.g. "BMW Group – Trainee Programme")
          company_addr  : str  (e.g. "BMW AG, Munich, Germany")
          subject_line  : str  (e.g. "Application – Software Engineer | Ref 123")
          para1         : str  (opening paragraph ~120 words)
          para2         : str  (experience paragraph ~120 words)
          para3         : str  (formula student paragraph ~80 words)
          para4         : str  (contribution paragraph ~80 words)
          para5         : str  (closing paragraph ~60 words)
        """
        doc = Document(str(template_path))
        paras = doc.paragraphs

        for section, indices in CL_SECTIONS.items():
            value = content.get(section, "")
            if not value:
                continue
            if indices[0] < len(paras):
                if section == "subject_line":
                    self._replace_subject_line(paras[indices[0]], value.strip())
                elif section in ("company_name", "company_addr"):
                    # Bold header lines — preserve bold=True, just swap text
                    self._replace_bold_header(paras[indices[0]], value.strip())
                else:
                    self._replace_cl_paragraph(paras[indices[0]], value.strip())

        # Update date in para 2 (tab + date + newline + "Recruiting Team")
        self._update_cl_date(paras[2] if len(paras) > 2 else None)

        doc.save(str(output_path))
        logger.info(f"CL saved → {output_path.name}")
        return output_path

    # ── Low-level replacers ────────────────────────────────────

    def _replace_bullet(self, para, text: str) -> None:
        """
        Replace a bullet paragraph that follows the format:
        'Bold Label: Regular description text.'

        The colon+space are part of the bold label.
        Bold=True on label run, bold=None (inherit) on body run — matches template.
        Strips any ** markers the AI may have added around the label.
        """
        # Strip ** markers and <b>...</b> HTML tags the AI sometimes adds
        text = re.sub(r'\*\*', '', text)
        text = re.sub(r'</?b>', '', text, flags=re.IGNORECASE)

        if ': ' in text:
            label_part, body_part = text.split(': ', 1)
            bold_text = label_part + ': '
            normal_text = body_part
        else:
            bold_text = ""
            normal_text = text

        ref_run = para.runs[0] if para.runs else None

        for run in para.runs:
            run.text = ""

        if para.runs:
            r0 = para.runs[0]
            r0.text = bold_text
            r0.bold = True

            if len(para.runs) > 1:
                for r in para.runs[1:]:
                    r.text = ""
                para.runs[1].text = normal_text
                para.runs[1].bold = None   # inherit — matches reference template
            else:
                new_run = para.add_run(normal_text)
                new_run.bold = None        # inherit — never explicitly False
                if ref_run:
                    try:
                        new_run.font.size = ref_run.font.size
                        new_run.font.name = ref_run.font.name
                    except Exception:
                        pass
        else:
            r_bold = para.add_run(bold_text)
            r_bold.bold = True
            r_norm = para.add_run(normal_text)
            r_norm.bold = None

    def _replace_plain(self, para, text: str) -> None:
        """Replace a paragraph with plain text (summary, competencies).
        Never forces bold=False — lets the paragraph style inherit naturally.
        """
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = text
            # Do NOT set bold — leave it as None so style is inherited
        else:
            para.add_run(text)

    def _replace_formula_student(self, para, text: str) -> None:
        """
        Replace Formula Student description paragraphs (Schanzer / Veloce).
        Template varies:
          - Para 38: run[0]='Description'(bold), run[1]=':'(bold), run[2]=' ', run[3+]=content
          - Para 40: run[0]='Description'(bold), run[1]=': ',                  run[2+]=content
        Auto-detects the first content run (first run whose stripped text is not
        a colon/space label fragment), then clears all content runs and writes there.
        """
        # Strip any ** markers — Formula Student descriptions are plain prose
        text = re.sub(r'\*\*', '', text)
        runs = para.runs
        if not runs:
            para.add_run(text)
            return

        # Find where content starts: first run after run[0] whose text is not
        # just a colon or whitespace (i.e., it is actual description text)
        content_start = len(runs)  # fallback: append
        for i in range(1, len(runs)):
            stripped = runs[i].text.strip().lstrip(':').strip()
            if stripped:  # non-empty after stripping colons/spaces → content run
                content_start = i
                break

        if content_start < len(runs):
            for run in runs[content_start:]:
                run.text = ""
            runs[content_start].text = text
        else:
            new_run = para.add_run(text)
            new_run.bold = None

    def _replace_bold_header(self, para, text: str) -> None:
        """
        Replace a single-run bold header line (company_name, company_addr).
        Keeps the existing bold=True state — only changes the text.
        """
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = text
            # Do NOT override bold — template already has bold=True
        else:
            r = para.add_run(text)
            r.bold = True

    def _replace_subject_line(self, para, text: str) -> None:
        """
        Replace subject line value while keeping the 'Subject' bold prefix.
        Template: run[0]='Subject'(bold), run[1]=': '(normal), run[2+]=value(normal).
        Only the value portion (runs 2+) is replaced.
        """
        if len(para.runs) >= 3:
            for run in para.runs[2:]:
                run.text = ""
            para.runs[2].text = text
        elif len(para.runs) == 2:
            # Subject + ': value' in two runs
            para.runs[1].text = ": " + text
        elif para.runs:
            para.runs[0].text = text
        else:
            para.add_run(text)

    def _replace_cl_paragraph(self, para, text: str) -> None:
        """
        Replace a CL body paragraph.
        Handles inline **bold markers**: 'At **Infineon Technologies**, I...'
        Non-bold segments use bold=None (inherit) — never bold=False.
        """
        parts = self._split_bold_markers(text)
        for run in para.runs:
            run.text = ""

        existing_count = len(para.runs)
        run_idx = 0

        for is_bold, segment_text in parts:
            if run_idx < existing_count:
                r = para.runs[run_idx]
                r.text = segment_text
                r.bold = True if is_bold else None   # None = inherit, not explicit False
                run_idx += 1
            else:
                r = para.add_run(segment_text)
                r.bold = True if is_bold else None

        for remaining in para.runs[run_idx:existing_count]:
            remaining.text = ""

    @staticmethod
    def _split_bold_markers(text: str) -> List[Tuple[bool, str]]:
        """
        Split text on **bold** markers into (is_bold, text) tuples.
        E.g. 'At **Infineon**, I did X' → [(False,'At '), (True,'Infineon'), (False,', I did X')]
        """
        parts = []
        pattern = re.compile(r'\*\*(.*?)\*\*')
        last = 0
        for m in pattern.finditer(text):
            if m.start() > last:
                parts.append((False, text[last:m.start()]))
            parts.append((True, m.group(1)))
            last = m.end()
        if last < len(text):
            parts.append((False, text[last:]))
        if not parts:
            parts = [(False, text)]
        return parts

    def _fix_emf_icons(self, doc) -> None:
        """
        Replace EMF icon images in the CV header contact line with Unicode equivalents.

        Root cause: Word embeds the email/phone icons as .emf (Windows Enhanced
        Metafile) drawings. LibreOffice cannot render .emf files, so they disappear
        in the PDF output — visible in Word, invisible in PDF.

        Fix: detect any w:drawing that references a .emf relationship in paragraph 1
        (the contact line), remove the drawing, and insert a Unicode character in its
        place using Segoe UI Symbol font (renders correctly in both Word and LibreOffice).

        Icon order in para 1:  first EMF = email ✉,  second EMF = phone ✆
        """
        _A_NS  = "http://schemas.openxmlformats.org/drawingml/2006/main"
        _R_NS  = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        _ICONS = ["✉", "✆"]   # email, phone — in order of appearance

        # Build rId → icon char map for all .emf relationships
        emf_rids: list[str] = []
        for rel in doc.part.rels.values():
            if hasattr(rel, "target_ref") and rel.target_ref.endswith(".emf"):
                emf_rids.append(rel.rId)

        if not emf_rids:
            return  # no EMF images — nothing to fix

        para = doc.paragraphs[1]   # contact info line
        icon_idx = 0

        for run in para.runs:
            if icon_idx >= len(_ICONS):
                break
            drawing = run._r.find(qn("w:drawing"))
            if drawing is None:
                continue
            blip = drawing.find(".//{%s}blip" % _A_NS)
            if blip is None:
                continue
            rid = blip.get("{%s}embed" % _R_NS, "")
            if rid not in emf_rids:
                continue

            # Remove the drawing element
            run._r.remove(drawing)

            # Ensure the run has an rPr with Segoe UI Symbol font
            rPr = run._r.find(qn("w:rPr"))
            if rPr is None:
                rPr = OxmlElement("w:rPr")
                run._r.insert(0, rPr)
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.insert(0, rFonts)
            rFonts.set(qn("w:ascii"), "Segoe UI Symbol")
            rFonts.set(qn("w:hAnsi"), "Segoe UI Symbol")
            rFonts.set(qn("w:cs"),    "Segoe UI Symbol")

            # Insert the unicode icon as text
            t = OxmlElement("w:t")
            t.text = _ICONS[icon_idx]
            run._r.append(t)

            icon_idx += 1

    def _fix_formula_student_bullets(self, paras) -> None:
        """
        Remove the sub-bullet numPr from the Description paragraphs (indices 38 and 40).
        Word renders the level-1 list bullet correctly; LibreOffice renders it as 'o'
        in a box. Stripping numPr and setting explicit left indent fixes both renderers.
        """
        for idx in [38, 40]:
            if idx >= len(paras):
                continue
            p = paras[idx]._p
            pPr = p.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                p.insert(0, pPr)
            # Remove numPr (the sub-bullet definition)
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                pPr.remove(numPr)
            # Ensure explicit left indent so text stays visually indented
            ind = pPr.find(qn('w:ind'))
            if ind is None:
                ind = OxmlElement('w:ind')
                pPr.append(ind)
            ind.set(qn('w:left'), '720')
            ind.attrib.pop(qn('w:hanging'), None)

    def _update_date(self, paras) -> None:
        """Update the date in the last paragraph containing 'Date: DD.MM.YYYY'.
        The label ('Date: ') and value live in separate runs —
        only replace the run that contains the date value pattern.
        """
        today = datetime.now().strftime("%d.%m.%Y")
        date_pattern = re.compile(r'\d{2}\.\d{2}\.\d{4}')
        for para in reversed(paras):
            if "Date:" in para.text:
                for run in para.runs:
                    if date_pattern.search(run.text):
                        run.text = date_pattern.sub(today, run.text)
                        return  # only update the first date-value run found

    def _update_cl_date(self, para) -> None:
        """Update date in CL para 2 ('\tApril 6, 2026\nRecruiting Team')."""
        if not para:
            return
        now = datetime.now()
        today = now.strftime("%B") + " " + str(now.day) + ", " + str(now.year)

        for run in para.runs:
            # Look for a date pattern in the run text
            updated = re.sub(
                r'(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d+,\s+\d{4}',
                today,
                run.text,
            )
            if updated != run.text:
                run.text = updated
                return

    # ── Utility ────────────────────────────────────────────────

    def get_cv_placeholders(self, template_path: Path) -> Dict[str, str | List[str]]:
        """Return the current text of all variable sections (useful for logging/debug)."""
        doc = Document(str(template_path))
        paras = doc.paragraphs
        result = {}

        for section, indices in CV_SECTIONS.items():
            if len(indices) == 1:
                result[section] = paras[indices[0]].text if indices[0] < len(paras) else ""
            else:
                result[section] = [
                    paras[i].text for i in indices if i < len(paras)
                ]
        return result

    def get_cl_placeholders(self, template_path: Path) -> Dict[str, str]:
        doc = Document(str(template_path))
        paras = doc.paragraphs
        return {
            section: paras[indices[0]].text if indices[0] < len(paras) else ""
            for section, indices in CL_SECTIONS.items()
        }
